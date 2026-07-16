#!/usr/bin/env python3
"""
Generic markdown → DOCX + PDF renderer with brand styling (no logo).

Usage:
    python scripts/render-docs.py references/operator-guide/01-quick-start.md
    python scripts/render-docs.py --all-operator-guides
    python scripts/render-docs.py --pdf-only references/operator-guide/01-quick-start.md
"""

import os
import sys
import subprocess
import tempfile
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Brand Colors (identical to _build_docx.py) ──────────────────────
SW_RED = RGBColor(0xCC, 0x41, 0x41)
SW_DARK = RGBColor(0x33, 0x33, 0x33)
SW_MED = RGBColor(0x79, 0x79, 0x79)
SW_ALT_ROW = 'F5F5F5'

# ── Styling helpers ───────────────────────────────────────────────────

def _add_heading(doc, text, level=1):
    """Brand-styled heading: red left-border on H1, bold Calibri."""
    p = doc.add_paragraph()
    if level == 1:
        # Red left border on H1
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        left = OxmlElement('w:left')
        left.set(qn('w:val'), 'single')
        left.set(qn('w:sz'), '8')
        left.set(qn('w:space'), '6')
        left.set(qn('w:color'), 'CC4141')
        pBdr.append(left)
        pPr.append(pBdr)
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(16)
        run.font.bold = True
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
    elif level == 2:
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(13)
        run.font.bold = True
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(2)
    else:
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.font.bold = True
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
    run.font.color.rgb = SW_DARK

def _add_body(doc, text):
    """Body paragraph with brand styling."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.color.rgb = SW_DARK
    p.paragraph_format.space_after = Pt(5)

def _add_code(doc, text):
    """Monospace code block with light grey background."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'F0F0F0')
    shading.set(qn('w:val'), 'clear')
    pPr.append(shading)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.5)

def _add_formatted_run(paragraph, text, font_name='Calibri', font_size=Pt(10), colour=SW_DARK, bold_default=False):
    """Add text to a paragraph, rendering **bold** markers as actual bold runs."""
    parts = text.split('**')
    for i, part in enumerate(parts):
        if not part:
            continue
        run = paragraph.add_run(part)
        run.font.name = font_name
        run.font.size = font_size
        run.font.color.rgb = colour
        # Odd-indexed parts (1,3,5...) are inside ** markers
        run.font.bold = bold_default or (i % 2 == 1)

def _add_table(doc, headers, rows):
    """Brand-styled table: white-on-dark header, alternating row shading."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        if '<br>' in h:
            parts = h.split('<br>')
            for pi, part in enumerate(parts):
                if pi > 0:
                    run_br = cell.paragraphs[0].add_run()
                    run_br.add_break()
                part = part.strip()
                if not part: continue
                run = cell.paragraphs[0].add_run(part)
                run.font.name = 'Calibri'
                run.font.size = Pt(8) if pi > 0 else Pt(8)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        else:
            run = cell.paragraphs[0].add_run(h)
            run.font.name = 'Calibri'
            run.font.size = Pt(8)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), '333333')
        shading.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shading)
    # Data rows
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = ''
            text = str(val)
            # Handle <br> line breaks in header cells
            if '<br>' in text:
                parts = text.split('<br>')
                for pi, part in enumerate(parts):
                    if pi > 0:
                        # Add line break
                        run_br = cell.paragraphs[0].add_run()
                        run_br.add_break()
                    part = part.strip()
                    if not part:
                        continue
                    # Colour emoji
                    if part.startswith('✅') or part.startswith('✔'):
                        run = cell.paragraphs[0].add_run(part)
                        run.font.name = 'Calibri'; run.font.size = Pt(8)
                        run.font.color.rgb = RGBColor(0x2D, 0xA4, 0x4E)
                    elif part.startswith('❌') or part.startswith('✘'):
                        run = cell.paragraphs[0].add_run(part)
                        run.font.name = 'Calibri'; run.font.size = Pt(8)
                        run.font.color.rgb = RGBColor(0xCC, 0x41, 0x41)
                    else:
                        run = cell.paragraphs[0].add_run(part)
                        run.font.name = 'Calibri'; run.font.size = Pt(7)
                        run.font.color.rgb = SW_MED if r == 0 else SW_DARK
                continue  # Skip normal rendering for <br> cells
            text = str(val)
            # Split emoji from text — colour emoji, keep text default
            if text.startswith('✅') or text.startswith('✔'):
                emoji_end = 1
                emoji_colour = RGBColor(0x2D, 0xA4, 0x4E)  # Green
            elif text.startswith('❌') or text.startswith('✘'):
                emoji_end = 1
                emoji_colour = RGBColor(0xCC, 0x41, 0x41)  # Red
            elif text.startswith('⚠️'):
                emoji_end = 2 if len(text) > 1 and text[1] == '️' else 1
                emoji_colour = RGBColor(0xE6, 0xA8, 0x17)  # Amber
            elif text.startswith('⚠'):
                emoji_end = 1
                emoji_colour = RGBColor(0xE6, 0xA8, 0x17)
            else:
                emoji_end = 0
                emoji_colour = None

            if emoji_end > 0:
                # Emoji part
                run = cell.paragraphs[0].add_run(text[:emoji_end])
                run.font.name = 'Calibri'; run.font.size = Pt(8)
                run.font.color.rgb = emoji_colour
                # Text part (default black) — with bold support
                rest = text[emoji_end:].strip()
                if rest:
                    _add_formatted_run(cell.paragraphs[0], ' ' + rest, font_size=Pt(8), colour=SW_DARK)
            else:
                _add_formatted_run(cell.paragraphs[0], text, font_size=Pt(8), colour=SW_DARK)
        # Alternating row shading
        if r % 2 == 0:
            for c in range(len(headers)):
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), SW_ALT_ROW)
                shading.set(qn('w:val'), 'clear')
                table.rows[r + 1].cells[c]._tc.get_or_add_tcPr().append(shading)
    # Set column widths — first column wider for label text
    if len(headers) > 4:
        first_w = Cm(3.5)  # Feature column
        rest_w = Cm(2.2)   # Other columns
        for i in range(len(headers)):
            w = first_w if i == 0 else rest_w
            for row in table.rows:
                row.cells[i].width = w
    doc.add_paragraph()

def _add_bullet(doc, text, indent_level=0):
    """Bullet point with brand styling."""
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.color.rgb = SW_DARK

def _add_checkbox(doc, text, checked=False):
    """Checkbox item: [x] or [ ] prefix."""
    prefix = '[x]' if checked else '[ ]'
    p = doc.add_paragraph()
    run = p.add_run(f'{prefix}  {text}')
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.color.rgb = SW_DARK
    p.paragraph_format.left_indent = Cm(0.5)

def _add_callout(doc, text):
    """Red background callout box."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    pPr = p._p.get_or_add_pPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'CC4141')
    shading.set(qn('w:val'), 'clear')
    pPr.append(shading)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


# ── Markdown parser ──────────────────────────────────────────────────

def _parse_markdown_simple(text):
    """
    Simple markdown → structured tokens.
    Handles: headings, paragraphs, code blocks (```), tables (|), bullets (-), checkboxes (- [ ])
    """
    lines = text.split('\n')
    tokens = []
    i = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_headers = []
    table_rows = []

    while i < len(lines):
        line = lines[i]

        # Code block toggle
        if line.strip().startswith('```'):
            if in_code_block:
                tokens.append(('code', '\n'.join(code_lines)))
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Empty line
        if not line.strip():
            # Flush pending table
            if in_table:
                tokens.append(('table', table_headers, table_rows))
                in_table = False
                table_headers = []
                table_rows = []
            i += 1
            continue

        # Heading
        if line.startswith('# ') or line.startswith('## ') or line.startswith('### ') or line.startswith('#### '):
            if in_table:
                tokens.append(('table', table_headers, table_rows))
                in_table = False
                table_headers = []
                table_rows = []
            level = len(line) - len(line.lstrip('#'))
            heading = line.lstrip('#').strip()
            tokens.append(('heading', level, heading))
            i += 1
            continue

        # Table row
        if '|' in line and line.strip().startswith('|'):
            cells = [c.strip() for c in line.strip().strip('|').split('|')]
            # Skip separator rows like |---|---|
            if all(c.replace('-', '').replace(':', '').strip() == '' for c in cells):
                i += 1
                continue
            if not in_table:
                table_headers = cells
                in_table = True
            else:
                table_rows.append(cells)
            i += 1
            continue

        # Checkbox
        if line.strip().startswith('- [ ]') or line.strip().startswith('- [x]') or line.strip().startswith('- [X]'):
            if in_table:
                tokens.append(('table', table_headers, table_rows))
                in_table = False
                table_headers = []
                table_rows = []
            checked = 'x' in line.strip()[3:5].lower()
            text = line.strip()[5:].strip()
            tokens.append(('checkbox', checked, text))
            i += 1
            continue

        # Bullet
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            if in_table:
                tokens.append(('table', table_headers, table_rows))
                in_table = False
                table_headers = []
                table_rows = []
            text = line.strip()[2:].strip()
            tokens.append(('bullet', text))
            i += 1
            continue

        # Callout (starts with >)
        if line.strip().startswith('>'):
            text = line.strip()[1:].strip()
            # Collect multi-line callouts
            callout_lines = [text]
            while i + 1 < len(lines) and lines[i + 1].strip().startswith('>'):
                i += 1
                callout_lines.append(lines[i].strip()[1:].strip())
            tokens.append(('callout', ' '.join(callout_lines)))
            i += 1
            continue

        # Regular paragraph
        if in_table:
            tokens.append(('table', table_headers, table_rows))
            in_table = False
            table_headers = []
            table_rows = []
        # Accumulate consecutive body lines
        body_lines = [line.strip()]
        while i + 1 < len(lines) and lines[i + 1].strip() and \
              not lines[i + 1].startswith('#') and not lines[i + 1].startswith('```') and \
              not lines[i + 1].startswith('|') and not lines[i + 1].strip().startswith('- ') and \
              not lines[i + 1].strip().startswith('>'):
            i += 1
            body_lines.append(lines[i].strip())
        tokens.append(('body', ' '.join(body_lines)))
        i += 1

    # Flush final table
    if in_table:
        tokens.append(('table', table_headers, table_rows))

    return tokens


# ── Document builder ──────────────────────────────────────────────────

def _setup_document(title):
    """Create a new DOCX with brand margins, header, and footer. No logo."""
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(1.5)
        s.bottom_margin = Cm(1.5)
        s.left_margin = Cm(1.8)
        s.right_margin = Cm(1.8)

    # Header — no logo, just title + separator line
    header = doc.sections[0].header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = hp.add_run(f'PORTSHIM  |  {title}')
    run.font.name = 'Calibri'
    run.font.size = Pt(9)
    run.font.color.rgb = SW_MED

    # Red bottom border on header
    p2 = header.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    pPr = p2._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), 'CC4141')
    pBdr.append(bottom)
    pPr.append(pBdr)

    # Footer
    footer = doc.sections[0].footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = fp.add_run('PortShim — Confidential')
    run.font.name = 'Calibri'
    run.font.size = Pt(8)
    run.font.color.rgb = SW_MED

    return doc


def _tokens_to_docx(doc, tokens):
    """Convert parsed markdown tokens to DOCX elements."""
    for token in tokens:
        kind = token[0]

        if kind == 'heading':
            _, level, text = token
            # Normalize level > 2 to 2 (H3+ treated as H2 style)
            _add_heading(doc, text, min(level, 2))

        elif kind == 'body':
            _add_body(doc, token[1])

        elif kind == 'code':
            for line in token[1].split('\n'):
                _add_code(doc, line)

        elif kind == 'table':
            _, headers, rows = token
            _add_table(doc, headers, rows)

        elif kind == 'bullet':
            _add_bullet(doc, token[1])

        elif kind == 'checkbox':
            _, checked, text = token
            _add_checkbox(doc, text, checked)

        elif kind == 'callout':
            _add_callout(doc, token[1])


def render_markdown(md_path, output_dir=None):
    """
    Convert a markdown file to DOCX and PDF with brand styling.

    Returns: (docx_path, pdf_path)
    """
    md_path = Path(md_path).resolve()
    if not md_path.exists():
        raise FileNotFoundError(f'Markdown file not found: {md_path}')

    if output_dir is None:
        # Default: mirrors references/ → output/
        output_dir = Path(str(md_path.parent).replace('references', 'output', 1))
        output_dir.mkdir(parents=True, exist_ok=True)
    else:
        output_dir = Path(output_dir)

    # Read markdown
    with open(md_path, 'r', encoding='utf-8') as f:
        md_text = f.read()

    # Extract title from first H1
    title = md_path.stem.replace('-', ' ').title()
    for line in md_text.split('\n'):
        if line.startswith('# '):
            title = line[2:].strip()
            break

    # Parse and build DOCX
    doc = _setup_document(title)
    tokens = _parse_markdown_simple(md_text)
    _tokens_to_docx(doc, tokens)

    # Save DOCX
    docx_path = output_dir / f'{md_path.stem}.docx'
    doc.save(str(docx_path))
    print(f'  DOCX: {docx_path} ({os.path.getsize(docx_path)} bytes)')

    # Convert to PDF via docx2pdf (Word COM)
    pdf_path = output_dir / f'{md_path.stem}.pdf'
    try:
        from docx2pdf import convert
        convert(str(docx_path), str(pdf_path))
        print(f'  PDF:  {pdf_path} ({os.path.getsize(pdf_path)} bytes)')
    except Exception as e:
        print(f'  PDF:  SKIPPED — {e}')
        pdf_path = None

    return docx_path, pdf_path


# ── CLI ───────────────────────────────────────────────────────────────

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(1)

    arg = sys.argv[1]

    if arg == '--all-operator-guides':
        guide_dir = Path(__file__).resolve().parent.parent / 'references' / 'operator-guide'
        md_files = sorted(guide_dir.glob('*.md'))
        if not md_files:
            print(f'No .md files found in {guide_dir}')
            sys.exit(1)
        print(f'Rendering {len(md_files)} operator guide(s)...\n')
        for md in md_files:
            print(f'  {md.name}:')
            render_markdown(str(md))
            print()
    else:
        md_path = Path(arg)
        render_markdown(str(md_path))
