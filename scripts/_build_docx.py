#!/usr/bin/env python3
"""Build brand-styled executive summary DOCX — auto-generated, not manually edited."""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SW_RED = RGBColor(0xCC, 0x41, 0x41)
SW_DARK = RGBColor(0x33, 0x33, 0x33)
SW_MED = RGBColor(0x79, 0x79, 0x79)

def _h(doc, text, level=1):
    p = doc.add_paragraph()
    if level == 1:
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        left = OxmlElement('w:left')
        left.set(qn('w:val'), 'single'); left.set(qn('w:sz'), '8')
        left.set(qn('w:space'), '6'); left.set(qn('w:color'), 'CC4141')
        pBdr.append(left); pPr.append(pBdr)
        run = p.add_run(text)
        run.font.name = 'Calibri'; run.font.size = Pt(16); run.font.bold = True
        p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(4)
    else:
        run = p.add_run(text)
        run.font.name = 'Calibri'; run.font.size = Pt(13); run.font.bold = True
        p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(2)
    run.font.color.rgb = SW_DARK

def _b(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text); run.font.name = 'Calibri'; run.font.size = Pt(10)
    run.font.color.rgb = SW_DARK; p.paragraph_format.space_after = Pt(5)

def _t(doc, headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT; table.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]; cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.font.name = 'Calibri'; run.font.size = Pt(9); run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), '333333'); shading.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shading)
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r+1].cells[c]; cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            run.font.name = 'Calibri'; run.font.size = Pt(9); run.font.color.rgb = SW_DARK
        # Alternating row shading
        if r % 2 == 0:
            for c in range(len(headers)):
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), 'F5F5F5'); shading.set(qn('w:val'), 'clear')
                table.rows[r+1].cells[c]._tc.get_or_add_tcPr().append(shading)
    doc.add_paragraph()

def _d(doc, term, desc):
    p = doc.add_paragraph()
    run = p.add_run(term); run.font.name = 'Calibri'; run.font.size = Pt(10)
    run.font.bold = True; run.font.color.rgb = SW_RED
    run = p.add_run(f'  {desc}'); run.font.name = 'Calibri'
    run.font.size = Pt(10); run.font.color.rgb = SW_DARK
    p.paragraph_format.space_after = Pt(2)

def _cta(doc, title, subtitle):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(14)
    pPr = p._p.get_or_add_pPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'CC4141'); shading.set(qn('w:val'), 'clear')
    pPr.append(shading)
    run = p.add_run(title); run.font.name = 'Calibri'; run.font.size = Pt(11)
    run.font.bold = True; run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    p2 = doc.add_paragraph()
    p2Pr = p2._p.get_or_add_pPr()
    shading2 = OxmlElement('w:shd')
    shading2.set(qn('w:fill'), 'CC4141'); shading2.set(qn('w:val'), 'clear')
    p2Pr.append(shading2)
    run2 = p2.add_run(subtitle); run2.font.name = 'Calibri'; run2.font.size = Pt(9)
    run2.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)

import os as _os
doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2); s.bottom_margin = Cm(2.2)
    s.left_margin = Cm(2.2); s.right_margin = Cm(2.2)

# Header — no logo
header = doc.sections[0].header; header.is_linked_to_previous = False
hp = header.paragraphs[0]; hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = hp.add_run("PORTSHIM  |  Executive Summary")
run.font.name = 'Calibri'; run.font.size = Pt(9); run.font.color.rgb = SW_MED
p2 = header.add_paragraph(); p2.paragraph_format.space_after = Pt(0)
pPr = p2._p.get_or_add_pPr(); pBdr = OxmlElement('w:pBdr')
bottom = OxmlElement('w:bottom')
bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '12')
bottom.set(qn('w:space'), '4'); bottom.set(qn('w:color'), 'CC4141')
pBdr.append(bottom); pPr.append(pBdr)

# Footer
footer = doc.sections[0].footer; footer.is_linked_to_previous = False
fp = footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = fp.add_run("PortShim - Confidential")
run.font.name = 'Calibri'; run.font.size = Pt(8); run.font.color.rgb = SW_MED

# Content (no body logo)
_h(doc, 'Executive Summary')
_b(doc, 'PortShim is a Hermes-orchestrated, distro-agnostic security assessment pipeline that maps a target network, cross-references known CVEs against discovered services, exploits vulnerabilities with IDS-aware stealth profiles, and produces a comprehensive report with an Excel remediation checklist that tracks fixes across retests.')
_b(doc, 'Designed for authorised on-site assessments, PortShim adapts to wired and wireless networks, Windows, Linux, and macOS targets, and operates in three LLM deployment modes: fully local (air-gapped), hybrid (tactical local + strategic cloud), and fully cloud.')
_t(doc, ['Metric', 'Value'], [['Total Components','40 files'],['Pipeline Phases','5 (Recon-Vuln-Exploit-Report-Retest)'],['Stealth Profiles','3 (Silent Entry, Surgical, Full Assault)']])

_h(doc, 'Pipeline Overview')
_h(doc, 'Phase 1 - Reconnaissance & Topology', 2)
_b(doc, 'Network discovery via nmap, httpx, subfinder, plus wireless (Kismet/Aircrack-ng). Structured host table with service fingerprinting and optional network topology diagram. Device classifier identifies 40+ device roles including firewalls, domain controllers, NVRs, IoT, hypervisors, and backup systems.')
_h(doc, 'Phase 2 - Vulnerability Analysis', 2)
_b(doc, 'Cross-references every discovered service version against the NVD via nmap-vulners. Nuclei deep-scans for known CVEs. Guardian Analyst correlates findings with CVSS v3.1 scoring.')
_h(doc, 'Phase 3 - Exploitation & Access', 2)
_b(doc, 'NeuroSploit (348 agents) and Guardian attack chains execute verified exploits. Anthropic skills cover Linux privilege escalation, Windows lateral movement, cloud exploitation (AWS/Azure/GCP), and Kubernetes penetration testing.')
_h(doc, 'Phase 4 - Reporting', 2)
_b(doc, 'Generates three deliverables: PowerPoint executive brief, Word technical report with CVSS methodology, and Excel remediation checklist with colour-coded severity and tick columns.')
_h(doc, 'Phase 5 - Retest & Verification', 2)
_b(doc, 'Compares baseline and retest scans. Auto-classifies findings as FIXED, STILL OPEN, NEW, or REGRESSION. Updates checklist with ticks and highlights new findings.')

_h(doc, 'Architecture')
_b(doc, 'PortShim is built on a five-layer stack with a cross-cutting knowledge-source tracking system that pins every external reference (GitHub repos, NVD feeds, tool docs) to specific commits and detects staleness.')
_d(doc, 'L5 - Orchestrator', 'site-assessment-pipeline skill drives the full 5-phase flow.')
_d(doc, 'L4 - Glue Scripts', '8 Python scripts: stealth profiles, topology map, device classifier, report generator (.docx/.pptx/.xlsx), Excel checklist, retest diff, LLM config, deployment bootstrap.')
_d(doc, 'L3 - Tools', '25+ Anthropic cybersecurity skills via npx + Guardian CLI + NeuroSploit + nmap-vulners NSE + nuclei + httpx + subfinder.')
_d(doc, 'L2 - Skills', 'Updated guardian-cli, neurosploit-integration, remote-system-access with knowledge-source tracking.')
_d(doc, 'L1 - Sources', 'sources.yaml per skill + sync_knowledge.py + check-skill-freshness.py. Commit-pinned references from GitHub repos.')

_h(doc, 'Built on Open Source')
_b(doc, 'PortShim leverages and extends the following community projects:')
_t(doc, ['Project','Stars','Role in PortShim'], [
    ['Guardian CLI','-','Primary scanning engine - 50+ security tools, multi-agent pipeline'],
    ['NeuroSploit','-','Autonomous exploitation - 348 agents, cross-model voting, attack chains'],
    ['nmap-vulners','3.3k','CVE cross-referencing - maps service versions to known vulnerabilities'],
    ['Anthropic Cybersecurity Skills','23.3k','25 targeted skills for wireless, privesc, lateral, cloud, reporting'],
    ['Red Teaming Toolkit','-','Curated tool reference organised by kill-chain phase'],
    ['nmap','-','Network discovery and port scanning'],
    ['nuclei','-','Template-based vulnerability scanning'],
    ['httpx','-','HTTP probing and tech stack detection'],
    ['subfinder','-','Subdomain enumeration'],
    ['Pwndoc','2.8k','Pentest report generation (reference)'],
    ['VECTR','1.6k','Red/blue team test tracking (reference)'],
    ['Vulnreport','600','Pentest management and automation (reference)'],
])
_b(doc, 'PortShim itself - the orchestrator, device classifier, stealth profiles, topology mapper, deployment bootstrap, and knowledge-source tracking - is original work. See skills/site-assessment-pipeline/sources.yaml for the complete dependency manifest with commit-pinned tracking.')

_h(doc, 'End-to-End Example: The Flat Network')
_b(doc, 'A representative scenario demonstrating what PortShim discovers during a typical on-site assessment of a medium-sized enterprise with a flat network architecture - no VLAN segmentation, no NAC, everything on one subnet.')
_h(doc, '08:00 - Arrival & Deployment', 2)
_b(doc, 'A single Linux laptop is connected to a wall port in a vacant meeting room. deploy.py detects Ubuntu 24.04, installs the full toolkit, pulls 25 Anthropic cybersecurity skills, and symlinks the orchestrator into Hermes. The engagement begins in Surgical mode.')
_h(doc, '08:04 - Reconnaissance & Device Classification', 2)
_b(doc, 'A /22 subnet sweep returns 847 live hosts. topology.py maps them. device-classifier.py immediately flags 23 high-value targets:')
_t(doc, ['Host','Role','Finding'], [['10.0.1.2','Domain Controller','LDAP without signing enforcement'],['10.0.1.10','Exchange 2016 CU22','ProxyShell (CVE-2021-34473)'],['10.0.1.50','MSSQL Server','Port 1433 exposed, xp_cmdshell enabled'],['10.0.2.5','UniFi Controller','CVE-2021-44529 - auth bypass'],['10.0.4.20','Hyper-V Host','WinRM on 5985 with NTLM, domain-joined']])
_h(doc, '08:18 - Exploitation', 2)
_b(doc, 'NeuroSploit cracks the UniFi controller hash within 90 seconds. SSH keys extracted from the controller give root access to all 47 campus APs. From the server room ceiling, cleartext LDAP authentication is captured. LLMNR poisoning yields a service account hash. That account has local admin on the MSSQL server. Mimikatz extracts the KRBTGT hash.')
_b(doc, 'Time from wall jack to domain dominance: 27 minutes. No zero-days used. Every exploit was a publicly known CVE with an available patch.')

_h(doc, 'Stealth Profiles')
_t(doc, ['Profile','Description','Nmap'], [['Silent Entry','IDS-aware, single-thread, 30s delay, no brute','T1 / -sT'],['Surgical','Rate-limited SYN, targeted nuclei, verified-CVE-only','T3 / -sS'],['Full Assault','Multi-threaded, full nuclei, all agents, parallel','T5 / -A']])

_h(doc, 'LLM Deployment Modes')
_t(doc, ['Mode','Description'], [['Fully Local','Everything on field laptop GPU. Uncensored local models. Air-gap capable.'],['Hybrid','Tactical scanning + exploitation stay local. Reports use cloud.'],['Fully Cloud','Thin client runs CLI tools only. All LLM work via API.']])

_h(doc, 'Testing Strategy')
_b(doc, 'PortShim ships with 155 unit tests across 8 test modules and 3 integration test modules. All unit tests run offline with mocked network calls.')

_h(doc, 'Deliverables')
_t(doc, ['Deliverable','Format','Description'], [['Executive Brief','.pptx','High-level findings, risk summary, recommended actions'],['Technical Report','.docx','Full finding details, CVSS vectors, remediation steps'],['Remediation Checklist','.xlsx','Color-coded, tick column, retest date tracking'],['Network Topology','.dot / table','Structured host table with CVE annotations'],['Retest Delta Report','.docx / .xlsx','Before/after: FIXED, STILL OPEN, NEW, REGRESSION']])

_h(doc, 'Deployment')
_b(doc, 'A single bootstrap script (deploy.py) detects the target Linux distribution by probing /etc/*-release files and selects the correct package manager (apt, dnf, pacman, zypper, or apk). It installs system packages, Go-based security tools, Python dependencies, and the Anthropic skill library - all in one command.')

_h(doc, 'Key Metrics')
_t(doc, ['Metric','Value'], [['Total Components','40 files (skills, scripts, tests, references, assets)'],['Custom Code','~400 lines Python (8 scripts)'],['Existing Tools','Guardian CLI (50+ tools), NeuroSploit (348 agents), Anthropic (25+ skills)'],['Stealth Profiles','3 (Silent Entry, Surgical, Full Assault)'],['LLM Modes','3 (Local, Hybrid, Cloud)'],['Report Formats','5 (.pptx, .docx, .xlsx, .dot, plain table)'],['Target Platforms','Wired + wireless, Windows + Linux + macOS + cloud + OT/ICS'],['Unit Tests','155 passing across 8 modules'],['Integration Tests','3 modules (knowledge-sync live, pipeline smoke, skill consistency)']])

_cta(doc, 'PortShim is ready for authorised deployment.', 'Run deploy.py on the target Linux machine to bootstrap the full capability.')

out = _os.path.expanduser('~/projects/portshim/output/exec-summary/exec-summary.docx')
doc.save(out)
print(f'OK: {out} ({_os.path.getsize(out)} bytes)')
